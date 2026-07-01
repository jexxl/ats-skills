# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "python-docx>=1.1",
#   "PyYAML>=6.0",
# ]
# ///
from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import tempfile
import urllib.error
import urllib.request
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path

import yaml

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
PLANTUML_BLOCK_RE = re.compile(r"```plantuml\s*\n(.*?)\n```", re.DOTALL)
FENCED_CODE_BLOCK_RE = re.compile(r"```.*?```", re.DOTALL)
KROKI_BASE_URL = os.environ.get("KROKI_BASE_URL", "https://kroki.io").rstrip("/")
KROKI_TIMEOUT_SECONDS = float(os.environ.get("KROKI_TIMEOUT_SECONDS", "30"))
DIAGRAM_IMAGE_MAX_WIDTH_CM = float(os.environ.get("DIAGRAM_IMAGE_MAX_WIDTH_CM", "13.5"))
DIAGRAM_IMAGE_MAX_HEIGHT_CM = float(os.environ.get("DIAGRAM_IMAGE_MAX_HEIGHT_CM", "20"))
DIAGRAM_IMAGE_PX_PER_CM = float(os.environ.get("DIAGRAM_IMAGE_PX_PER_CM", str(96 / 2.54)))
DIAGRAM_IMAGE_SCALE = float(os.environ.get("DIAGRAM_IMAGE_SCALE", "1.2"))
FONT_NAME = "仿宋"
CODE_FONT_NAME = "Consolas"
CODE_STYLE_NAMES = {"Source Code", "Verbatim Char"}
CODE_STYLE_IDS = {"SourceCode", "VerbatimChar"}
PANDOC_MARKDOWN_FORMAT = (
    "markdown"
    "+yaml_metadata_block"
    "+pipe_tables"
    "+fenced_code_blocks"
    "+tex_math_dollars"
    "+tex_math_single_backslash"
    "+tex_math_double_backslash"
)
UNSUPPORTED_LATEX_PATTERNS = {
    r"\\newcommand": "不要使用自定义命令；请直接写展开后的公式。",
    r"\\renewcommand": "不要使用自定义命令；请直接写展开后的公式。",
    r"\\def": "不要使用自定义命令；请直接写展开后的公式。",
    r"\\DeclareMathOperator": "不要声明新算子；请使用 \\operatorname{...}。",
    r"\\label": "Word 交付文档不使用 LaTeX 交叉引用。",
    r"\\ref": "Word 交付文档不使用 LaTeX 交叉引用。",
    r"\\eqref": "Word 交付文档不使用 LaTeX 交叉引用。",
    r"\\tag": "不要使用手写公式编号；如需编号请在正文说明。",
    r"\\begin\{equation\*?\}": "不要使用 equation 环境；请改用 $$...$$。",
    r"\\begin\{align\*?\}": "不要使用 align 环境；请拆成多个 $$...$$ 公式块。",
    r"\\begin\{aligned\}": "不要使用 aligned 环境；请拆成多个 $$...$$ 公式块。",
    r"\\begin\{gather\*?\}": "不要使用 gather 环境；请拆成多个 $$...$$ 公式块。",
    r"\\begin\{multline\*?\}": "不要使用 multline 环境；请拆成多个 $$...$$ 公式块。",
    r"\\begin\{array\}": "不要使用 array 环境；请改写为正文说明或 Markdown 表格。",
    r"\\begin\{matrix\}": "不要使用 matrix 环境；请改写为正文说明或 Markdown 表格。",
    r"\\begin\{pmatrix\}": "不要使用 matrix 环境；请改写为正文说明或 Markdown 表格。",
    r"\\begin\{bmatrix\}": "不要使用 matrix 环境；请改写为正文说明或 Markdown 表格。",
}


@dataclass(frozen=True)
class DocumentJob:
    source: Path
    output: Path
    reference: Path
    layout: str = "markdown"


REPORT_SECTION_TITLES = ("测试数据与配置", "测试过程", "测试结果", "测试结论")
REPORT_SUBSECTIONS = {
    "测试数据与配置": ("测试输入文件清单", "实际输入数据", "可泛化参数"),
    "测试过程": ("运行方式", "输出文件清单"),
}
REQUIRED_METADATA_KEYS = (
    "课题名称",
    "专题名称",
    "模型编号",
    "模型名称",
    "模型功能描述",
    "输入数据概括",
    "输出数据概括",
    "模型服务场景",
    "上游接口模型编号",
    "下游接口模型编号",
    "交付时间",
    "责任单位",
)
UNRESOLVED_PLACEHOLDER_RE = re.compile(
    r"\{[^{}\n]*(?:课题|专题|算法|模型|输入|输出|文件|字段|参数|默认值|取值|说明|当前日期|交付|责任|单位|接口|上游|下游|场景|名称|真实|运行|环境|日志|状态|功能|来源|报告|数据|配置|结论)[^{}\n]*\}"
)
PLACEHOLDER_VALUE_RE = re.compile(r"^\{[^{}\n]+\}$")


def main() -> None:
    args = parse_args()
    config = build_config(args)

    temp_root = Path(tempfile.mkdtemp(prefix="ats-algorithm-docs-"))
    try:
        jobs = build_jobs(config)
        image_dir = temp_root / "images"
        image_dir.mkdir(parents=True, exist_ok=True)

        for job in jobs:
            print(f"[docs] build {job.output}")
            preprocessed_md = temp_root / f"{job.source.stem}.pandoc.md"
            preprocess_markdown(job.source, preprocessed_md, image_dir, job.layout)
            run_pandoc(preprocessed_md, job.output, job.reference)
            normalize_docx(job.output, job.layout)

        if not bool(config.get("keep_temp", False)):
            shutil.rmtree(temp_root, ignore_errors=True)
        else:
            print(f"[docs] temp files kept at {temp_root}")
    except Exception:
        shutil.rmtree(temp_root, ignore_errors=True)
        raise


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build ATS algorithm report Word documents from Markdown.")
    parser.add_argument("--docs-dir", type=Path, required=True, help="Directory containing the completed Markdown docs. Also used as output directory unless --output-dir is set.")
    parser.add_argument("--output-dir", type=Path, help="Directory for generated .docx files.")
    parser.add_argument("--reference-dir", type=Path, default=ROOT / "docx-reference", help="Directory containing reference .docx files.")
    parser.add_argument("--algorithm-id", required=True, help="Algorithm id, for example 4-4-Z-1.")
    parser.add_argument("--keep-temp", action="store_true", help="Keep preprocessed Markdown and rendered images for debugging.")
    return parser.parse_args()


def build_config(args: argparse.Namespace) -> dict:
    docs_dir = args.docs_dir.resolve()
    output_dir = (args.output_dir or args.docs_dir).resolve()
    reference_dir = args.reference_dir.resolve()
    report_md = f"{args.algorithm_id}_原始实验运行数据及测试报告.md"
    documents = [
        {
            "source": report_md,
            "reference": "experiment-report-reference.docx",
            "output": report_md.replace(".md", ".docx"),
            "layout": "experiment_report",
        },
    ]
    return {
        "template_dir": str(docs_dir),
        "reference_dir": str(reference_dir),
        "output_dir": str(output_dir),
        "keep_temp": args.keep_temp,
        "documents": documents,
    }


def resolve_path(config: dict, value: str) -> Path:
    raw = Path(value)
    if raw.is_absolute():
        return raw
    return (ROOT / raw).resolve()


def build_jobs(config: dict) -> list[DocumentJob]:
    template_dir = resolve_path(config, config["template_dir"])
    reference_dir = resolve_path(config, config["reference_dir"])
    output_dir = resolve_path(config, config["output_dir"])
    output_dir.mkdir(parents=True, exist_ok=True)

    jobs: list[DocumentJob] = []
    for item in config["documents"]:
        source_name = item["source"]
        reference_name = item.get("reference", source_name.replace(".md", ".docx"))
        output_name = item.get("output", source_name.replace(".md", ".docx"))
        job = DocumentJob(
            source=(template_dir / source_name).resolve(),
            reference=(reference_dir / reference_name).resolve(),
            output=(output_dir / output_name).resolve(),
            layout=item.get("layout", "markdown"),
        )
        for path, label in [(job.source, "Markdown template"), (job.reference, "reference docx")]:
            if not path.exists():
                raise FileNotFoundError(f"{label} not found: {path}")
        jobs.append(job)
    return jobs


def preprocess_markdown(source: Path, destination: Path, image_dir: Path, layout: str) -> None:
    text = source.read_text(encoding="utf-8")
    if layout == "experiment_report":
        text = render_experiment_report_markdown(text)
    elif layout != "markdown":
        raise ValueError(f"Unsupported document layout: {layout}")
    text = replace_packaging_placeholders(text)
    if layout == "experiment_report":
        validate_no_unresolved_placeholders(text, source)
    validate_formula_syntax(text, source)
    block_index = 0

    def replace_block(match: re.Match[str]) -> str:
        nonlocal block_index
        block_index += 1
        code = match.group(1).strip()
        svg_path = image_dir / f"{source.stem}-flowchart-{block_index}.svg"
        width, height = render_plantuml_diagram(code, svg_path)
        width_cm, height_cm = fit_image_size_cm(width, height)
        return f"![]({svg_path.as_posix()}){{width={width_cm:.2f}cm height={height_cm:.2f}cm}}"

    text = PLANTUML_BLOCK_RE.sub(replace_block, text)
    destination.write_text(text, encoding="utf-8")


def validate_formula_syntax(text: str, source: Path) -> None:
    searchable = FENCED_CODE_BLOCK_RE.sub("", text)
    violations: list[str] = []
    for pattern, guidance in UNSUPPORTED_LATEX_PATTERNS.items():
        if re.search(pattern, searchable):
            violations.append(f"- `{pattern}`: {guidance}")
    if violations:
        raise ValueError(
            f"{source} contains LaTeX syntax that is not supported for stable Word conversion.\n"
            "Use simple pandoc-compatible math instead:\n"
            + "\n".join(violations)
        )


def render_experiment_report_markdown(text: str) -> str:
    metadata, body = split_frontmatter(text)
    validate_required_metadata(metadata)
    algorithm_id = meta_value(metadata, "模型编号") or extract_algorithm_id_from_text(text)
    algorithm_name = meta_value(metadata, "模型名称")
    title = " ".join(part for part in [algorithm_id or first_heading(body), algorithm_name] if part).strip()
    if not title:
        title = "{算法编号} {算法名称}"

    intro = extract_section(body, "算法模型简介").strip()
    if not intro:
        raise ValueError("Experiment report must contain a non-empty `## 算法模型简介` section.")

    flowchart_title, flowchart_note, flowchart = extract_flowchart_section(body)
    if not flowchart:
        raise ValueError("Experiment report must contain a PlantUML flowchart section.")

    inputs = extract_data_requirements(body, "输入数据要求")
    outputs = extract_data_requirements(body, "输出数据要求")
    validate_data_requirements(inputs, "输入数据要求")
    validate_data_requirements(outputs, "输出数据要求")
    report_sections = extract_required_report_sections(body)

    lines = [
        f"# {title}",
        "",
        "## 算法模型基本信息",
        "",
    ]
    lines.extend(build_basic_info_table(metadata, inputs, outputs))
    lines.extend(
        [
            "",
            "## 算法模型简介",
            "",
            intro,
            "",
            f"## {flowchart_title}",
            "",
        ]
    )
    if flowchart_note:
        lines.extend([flowchart_note, ""])
    lines.extend(["```plantuml", flowchart, "```"])
    for section_title, section_body in report_sections:
        lines.extend(["", f"## {section_title}", "", section_body])
    return "\n".join(lines).rstrip() + "\n"


def extract_required_report_sections(text: str) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    for title in REPORT_SECTION_TITLES:
        section_body = extract_section(text, title).strip()
        if not section_body:
            raise ValueError(f"Experiment report must contain a non-empty `## {title}` section.")
        for subsection in REPORT_SUBSECTIONS.get(title, ()):
            if not extract_markdown_subsection(section_body, subsection).strip():
                raise ValueError(f"Experiment report `## {title}` must contain a non-empty `### {subsection}` subsection.")
        if title == "测试数据与配置":
            actual_input = extract_markdown_subsection(section_body, "实际输入数据")
            if not re.search(r"^####\s+.+", actual_input, re.MULTILINE):
                raise ValueError("Experiment report `### 实际输入数据` must contain at least one `#### 输入数据名称` item.")
        if title == "测试结果":
            if not re.search(r"^###\s+.+", section_body, re.MULTILINE):
                raise ValueError("Experiment report `## 测试结果` must contain at least one `### 输出数据名称` item.")
            if not (re.search(r"^\|.+\|$", section_body, re.MULTILINE) or re.search(r"```", section_body)):
                raise ValueError("Experiment report `## 测试结果` must include a real output table or fenced summary block.")
        sections.append((title, section_body))
    return sections


def extract_markdown_subsection(text: str, title: str) -> str:
    pattern = re.compile(rf"^###\s+{re.escape(title)}\s*$\n(.*?)(?=^###\s+|\Z)", re.MULTILINE | re.DOTALL)
    match = pattern.search(text)
    return match.group(1).strip() if match else ""


def validate_no_unresolved_placeholders(text: str, source: Path) -> None:
    matches = sorted(set(match.group(0) for match in UNRESOLVED_PLACEHOLDER_RE.finditer(text)))
    if matches:
        joined = "\n".join(f"- {match}" for match in matches[:20])
        raise ValueError(f"{source} contains unresolved report placeholders:\n{joined}")


def validate_required_metadata(metadata: dict) -> None:
    missing = [key for key in REQUIRED_METADATA_KEYS if is_missing_metadata_value(metadata, key)]
    if missing:
        joined = "\n".join(f"- {key}" for key in missing)
        raise ValueError(
            "Experiment report frontmatter is missing required user-provided metadata:\n"
            f"{joined}\n"
            "Do not rely on defaults for ATS business metadata."
        )


def is_missing_metadata_value(metadata: dict, key: str) -> bool:
    value = meta_value(metadata, key)
    return not value or bool(PLACEHOLDER_VALUE_RE.fullmatch(value))


def validate_data_requirements(items: list[dict], section_title: str) -> None:
    if not items:
        raise ValueError(
            f"{section_title} section has no parseable data blocks. Use one line for the data name, "
            "the next line for `(`文件名`)`, then a Markdown table with headers: 字段中文 | 字段英文 | 字段类型."
        )
    invalid = [item for item in items if not item["fields"]]
    if invalid:
        names = "、".join(f"{item['name']}（{item['file']}）" for item in invalid)
        raise ValueError(
            f"{section_title} has data blocks without parseable field rows: {names}. "
            "The table header must be: | 字段中文 | 字段英文 | 字段类型 |."
        )


def split_frontmatter(text: str) -> tuple[dict, str]:
    match = re.match(r"^---\s*\n(.*?)\n---\s*\n?(.*)$", text, re.DOTALL)
    if not match:
        return {}, text
    metadata = yaml.safe_load(match.group(1)) or {}
    return metadata, match.group(2)


def replace_packaging_placeholders(text: str) -> str:
    metadata, _ = split_frontmatter(text)
    algorithm_id = clean_placeholder(metadata.get("模型编号")) or extract_algorithm_id_from_text(text)
    package_id = to_package_id(algorithm_id)
    if package_id:
        text = text.replace("{打包编号}", package_id)
    return text


def extract_algorithm_id_from_text(text: str) -> str:
    match = re.search(r"\b(\d+-\d+-[JjZz]-\d+)\b", text)
    return match.group(1) if match else ""


def to_package_id(algorithm_id: str) -> str:
    algorithm_id = clean_placeholder(algorithm_id)
    match = re.fullmatch(r"(\d+-\d+-[JjZz]-\d+)", algorithm_id)
    if not match:
        return ""
    return f"algo{match.group(1).lower()}"


def first_heading(text: str) -> str:
    match = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
    return match.group(1).strip() if match else ""


def extract_section(text: str, title: str) -> str:
    pattern = re.compile(rf"^##\s+{re.escape(title)}\s*$\n(.*?)(?=^##\s+|\Z)", re.MULTILINE | re.DOTALL)
    match = pattern.search(text)
    return match.group(1).strip() if match else ""


def extract_plantuml(text: str) -> str:
    match = PLANTUML_BLOCK_RE.search(text)
    return match.group(1).strip() if match else ""


def extract_flowchart_section(text: str) -> tuple[str, str, str]:
    sections = re.finditer(r"^##\s+(.+?)\s*$\n(.*?)(?=^##\s+|\Z)", text, re.MULTILINE | re.DOTALL)
    for section in sections:
        section_title = section.group(1).strip()
        section_body = section.group(2).strip()
        plantuml = extract_plantuml(section_body)
        if plantuml:
            note = PLANTUML_BLOCK_RE.sub("", section_body).strip()
            return section_title, note, plantuml
    return "算法模型流程图", "", ""


def extract_data_requirements(text: str, title: str) -> list[dict]:
    section = extract_subsection(text, title)
    blocks = re.split(r"\n(?=[^\n|（].*?\n\s*\(`[^`]+`\)\s*\n)", section)
    requirements: list[dict] = []
    for block in blocks:
        header = re.search(r"^\s*(.+?)\s*\n\s*\(`([^`]+)`\)\s*$", block, re.MULTILINE)
        if not header:
            continue
        requirements.append(
            {
                "name": header.group(1).strip(),
                "file": header.group(2).strip(),
                "fields": parse_markdown_table(block),
            }
        )
    return requirements


def extract_subsection(text: str, title: str) -> str:
    pattern = re.compile(rf"^###\s+{re.escape(title)}\s*$\n(.*?)(?=^###\s+|^##\s+|\Z)", re.MULTILINE | re.DOTALL)
    match = pattern.search(text)
    return match.group(1).strip() if match else ""


def parse_markdown_table(text: str) -> list[dict]:
    table_lines = [line.strip() for line in text.splitlines() if line.strip().startswith("|")]
    if len(table_lines) < 3:
        return []
    headers = split_table_row(table_lines[0])
    required_headers = {"字段中文", "字段英文", "字段类型"}
    if not required_headers.issubset(set(headers)):
        return []
    fields: list[dict] = []
    for line in table_lines[2:]:
        values = split_table_row(line)
        if len(values) != len(headers):
            continue
        row = dict(zip(headers, values))
        if not any(row.get(header, "").strip() for header in required_headers):
            continue
        fields.append(
            {
                "zh": row.get("字段中文", "").strip(),
                "en": row.get("字段英文", "").strip(),
                "type": row.get("字段类型", "").strip(),
            }
        )
    return fields


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def build_basic_info_table(metadata: dict, inputs: list[dict], outputs: list[dict]) -> list[str]:
    rows: list[list[str]] = [
        ["课题名称", meta_value(metadata, "课题名称")],
        ["专题名称", meta_value(metadata, "专题名称")],
        ["模型编号", meta_value(metadata, "模型编号")],
        ["模型名称", meta_value(metadata, "模型名称")],
        ["模型功能描述", meta_value(metadata, "模型功能描述")],
    ]
    rows.extend(requirement_rows("输入数据要求", metadata, "输入数据概括", inputs))
    rows.extend(requirement_rows("输出数据要求", metadata, "输出数据概括", outputs))
    rows.extend(
        [
            ["模型服务场景", meta_value(metadata, "模型服务场景")],
            ["上游接口模型编号", meta_value(metadata, "上游接口模型编号")],
            ["下游接口模型编号", meta_value(metadata, "下游接口模型编号")],
            ["交付时间", meta_value(metadata, "交付时间")],
            ["责任单位", meta_value(metadata, "责任单位")],
        ]
    )

    lines = ["| 项目 | 内容 | 字段 |", "| --- | --- | --- |"]
    for row in rows:
        padded = row + [""] * (3 - len(row))
        lines.append("| " + " | ".join(escape_table_cell(cell) for cell in padded[:3]) + " |")
    return lines


def requirement_rows(label: str, metadata: dict, summary_key: str, items: list[dict]) -> list[list[str]]:
    rows = [[label, meta_value(metadata, summary_key) or summarize_requirements(items), ""]]
    for item in items:
        fields = item["fields"] or [{"zh": "", "en": "", "type": ""}]
        first = True
        for field in fields:
            rows.append(
                [
                    "",
                    f"{item['name']}<br>`{item['file']}`" if first else "",
                    format_field(field),
                ]
            )
            first = False
    return rows


def summarize_requirements(items: list[dict]) -> str:
    if not items:
        return ""
    return "；".join(f"{item['name']}<br>`{item['file']}`" for item in items)


def format_field(field: dict) -> str:
    zh = field.get("zh", "")
    en = strip_inline_code(field.get("en", ""))
    field_type = strip_inline_code(field.get("type", ""))
    name = f"{zh}<br>`{en}`".strip() if en else zh
    return f"{name} (`{field_type}`)" if field_type else name


def strip_inline_code(value: str) -> str:
    value = str(value).strip()
    if len(value) >= 2 and value.startswith("`") and value.endswith("`"):
        return value[1:-1].strip()
    return value


def meta_value(metadata: dict, key: str, default: str = "") -> str:
    return clean_placeholder(metadata.get(key, default))


def clean_placeholder(value) -> str:
    if value is None:
        return ""
    if isinstance(value, dict) and len(value) == 1:
        key, dict_value = next(iter(value.items()))
        if dict_value is None:
            return "{" + str(key).strip() + "}"
    return str(value).replace("{ ", "{").replace(" }", "}").strip()


def escape_table_cell(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", "<br>")


def fit_image_size_cm(width: float, height: float) -> tuple[float, float]:
    width_cm = width / DIAGRAM_IMAGE_PX_PER_CM * DIAGRAM_IMAGE_SCALE
    height_cm = height / DIAGRAM_IMAGE_PX_PER_CM * DIAGRAM_IMAGE_SCALE
    shrink = min(
        DIAGRAM_IMAGE_MAX_WIDTH_CM / width_cm,
        DIAGRAM_IMAGE_MAX_HEIGHT_CM / height_cm,
        1.0,
    )
    width_cm *= shrink
    height_cm *= shrink
    return width_cm, height_cm


def read_svg_size(path: Path) -> tuple[float, float]:
    try:
        root = ET.fromstring(path.read_text(encoding="utf-8"))
    except ET.ParseError as exc:
        raise RuntimeError(f"Kroki did not return a valid SVG: {path}") from exc

    width = parse_svg_length(root.get("width", ""))
    height = parse_svg_length(root.get("height", ""))
    if width and height:
        return width, height

    view_box = root.get("viewBox", "")
    parts = view_box.replace(",", " ").split()
    if len(parts) == 4:
        try:
            return float(parts[2]), float(parts[3])
        except ValueError:
            pass
    raise RuntimeError(f"Cannot determine SVG size: {path}")


def parse_svg_length(value: str) -> float:
    match = re.match(r"^\s*([0-9.]+)", value)
    return float(match.group(1)) if match else 0.0


def render_plantuml_diagram(code: str, output: Path) -> tuple[float, float]:
    write_kroki_plantuml_svg(ensure_plantuml_document(code), output)
    return read_svg_size(output)


def ensure_plantuml_document(code: str) -> str:
    stripped = code.strip()
    if stripped.lower().startswith("@start"):
        return stripped
    return f"@startuml\n{stripped}\n@enduml"


def write_kroki_plantuml_svg(code: str, output: Path) -> None:
    fmt = output.suffix.lower().lstrip(".")
    if fmt != "svg":
        raise ValueError(f"Unsupported PlantUML output format: {output.suffix}")

    url = f"{KROKI_BASE_URL}/plantuml/{fmt}"
    request = urllib.request.Request(
        url,
        data=code.encode("utf-8"),
        headers={
            "Content-Type": "text/plain; charset=utf-8",
            "User-Agent": "ats-algorithm-packaging-docs/1.0",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=KROKI_TIMEOUT_SECONDS) as response:
            output.write_bytes(response.read())
    except urllib.error.HTTPError as exc:
        details = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"Kroki failed to render PlantUML as {fmt}: HTTP {exc.code}\n{details}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"Kroki request failed for {url}: {exc.reason}") from exc


def run_pandoc(source: Path, output: Path, reference: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    temp_output = output.with_name(f".{output.stem}.tmp{output.suffix}")
    command = [
        "pandoc",
        str(source),
        "--from",
        PANDOC_MARKDOWN_FORMAT,
        "--to",
        "docx",
        "--reference-doc",
        str(reference),
        "--output",
        str(temp_output),
    ]
    result = subprocess.run(command, cwd=ROOT, text=True, capture_output=True)
    if result.returncode != 0:
        raise RuntimeError(f"pandoc failed for {source}:\n{result.stderr.strip()}")
    try:
        temp_output.replace(output)
    except PermissionError as exc:
        raise PermissionError(f"Cannot replace {output}. Close the Word document or preview window and rerun.") from exc


def normalize_docx(path: Path, layout: str) -> None:
    doc = Document(path)
    normalize_sections(doc)
    normalize_styles(doc)
    normalize_markdown_heading_styles(doc)
    normalize_paragraphs(doc)
    normalize_tables(doc)
    if layout == "experiment_report":
        normalize_model_description(doc)
    doc.save(path)


def normalize_sections(doc: Document) -> None:
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)


def normalize_styles(doc: Document) -> None:
    style_sizes = {
        "Normal": 12,
        "Title": 22,
        "Heading 1": 16,
        "Heading 2": 12,
        "Heading 3": 12,
    }
    for style_name, size in style_sizes.items():
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = Pt(size)
        set_east_asia_font(style.element.rPr, FONT_NAME)
    for style_name in ["Source Code", "Verbatim Char"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = CODE_FONT_NAME
            style.font.size = Pt(10)
            set_east_asia_font(style.element.rPr, CODE_FONT_NAME)


def is_code_paragraph(paragraph) -> bool:
    p_style = paragraph._p.pPr.find(qn("w:pStyle")) if paragraph._p.pPr is not None else None
    return style_name(paragraph) == "Source Code" or (p_style is not None and p_style.get(qn("w:val")) in CODE_STYLE_IDS)


def is_code_run(run) -> bool:
    r_style = run._element.rPr.find(qn("w:rStyle")) if run._element.rPr is not None else None
    return style_name(run) in CODE_STYLE_NAMES or (r_style is not None and r_style.get(qn("w:val")) in CODE_STYLE_IDS)


def style_name(obj) -> str:
    return obj.style.name if obj.style else ""


def apply_run_font(run, font_name: str, size=None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = size
    set_east_asia_font(run._element.get_or_add_rPr(), font_name)


def add_line_breaks_before_inline_code(paragraph) -> None:
    preceding_text = ""
    for run in paragraph.runs:
        text = run.text or ""
        if is_code_run(run) and should_break_before_code(preceding_text, text):
            insert_break_at_run_start(run)
        preceding_text += text


def should_break_before_code(preceding_text: str, code_text: str) -> bool:
    if not preceding_text.strip() or not code_text.strip():
        return False
    previous_char = preceding_text.rstrip()[-1]
    return previous_char not in "([{（【:：/\\-"


def insert_break_at_run_start(run) -> None:
    if run._element.find(qn("w:br")) is not None:
        return
    br = OxmlElement("w:br")
    run._element.insert(1 if run._element.rPr is not None else 0, br)


def normalize_paragraphs(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        fmt = paragraph.paragraph_format
        if paragraph._p.findall(".//w:drawing", namespaces=paragraph._p.nsmap):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.style and paragraph.style.name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_after = Pt(12)
        elif paragraph.style and paragraph.style.name.startswith("Heading"):
            fmt.space_before = Pt(12)
            fmt.space_after = Pt(6)
        else:
            fmt.line_spacing = 1.0
            fmt.space_after = Pt(6)
        for run in paragraph.runs:
            if is_code_paragraph(paragraph) or is_code_run(run):
                apply_run_font(run, CODE_FONT_NAME)
            else:
                apply_run_font(run, FONT_NAME)


def normalize_markdown_heading_styles(doc: Document) -> None:
    has_title = any(paragraph.style and paragraph.style.name == "Title" for paragraph in doc.paragraphs)
    heading_map = (
        {
            "Heading 2": "Heading 1",
            "Heading 3": "Heading 2",
            "Heading 4": "Heading 3",
        }
        if has_title
        else {
            "Heading 1": "Title",
            "Heading 2": "Heading 1",
            "Heading 3": "Heading 2",
            "Heading 4": "Heading 3",
        }
    )
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name in heading_map and heading_map[style_name] in doc.styles:
            paragraph.style = doc.styles[heading_map[style_name]]


def normalize_tables(doc: Document) -> None:
    usable_width_twips = int((21.0 - 3.17 - 3.17) / 2.54 * 1440)
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        if "Table Grid" in doc.styles:
            table.style = "Table Grid"
        set_table_width(table, usable_width_twips)
        column_count = max(len(row.cells) for row in table.rows)
        column_width = int(usable_width_twips / column_count) if column_count else usable_width_twips
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_width(cell, column_width)
                set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        font_name = CODE_FONT_NAME if is_code_run(run) else FONT_NAME
                        apply_run_font(run, font_name, Pt(10.5 if row_index == 0 else 10))
                    add_line_breaks_before_inline_code(paragraph)


def normalize_model_description(doc: Document) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    delete_table_header_row(table)
    merge_model_description_table(table)
    apply_reference_table_format(table)
    remove_empty_image_captions(doc)
    add_page_break_before_first_drawing_heading(doc)
    add_page_break_after_first_drawing(doc)
    center_drawing_paragraphs(doc)


def delete_table_header_row(table) -> None:
    if not table.rows:
        return
    first_row_text = [cell.text.strip() for cell in table.rows[0].cells]
    if first_row_text[:3] == ["项目", "内容", "字段"]:
        table._tbl.remove(table.rows[0]._tr)


def merge_model_description_table(table) -> None:
    row_count = len(table.rows)
    if row_count == 0:
        return

    input_start = find_row_by_first_cell(table, "输入数据要求")
    output_start = find_row_by_first_cell(table, "输出数据要求")
    service_start = find_row_by_first_cell(table, "模型服务场景")

    simple_rows = set(range(0, min(input_start, row_count) if input_start != -1 else min(5, row_count)))
    if output_start != -1:
        simple_rows.add(output_start)
    if input_start != -1:
        simple_rows.add(input_start)
    for label in ["模型服务场景", "上游接口模型编号", "下游接口模型编号", "交付时间", "责任单位"]:
        index = find_row_by_first_cell(table, label)
        if index != -1:
            simple_rows.add(index)

    for row_index in sorted(simple_rows):
        if len(table.rows[row_index].cells) >= 3:
            safe_merge(table.cell(row_index, 1), table.cell(row_index, 2))

    if input_start != -1 and output_start != -1 and output_start > input_start:
        safe_merge(table.cell(input_start, 0), table.cell(output_start - 1, 0))
        merge_item_name_cells(table, input_start + 1, output_start)
    if output_start != -1:
        group_end = service_start if service_start != -1 else row_count
        if group_end > output_start:
            safe_merge(table.cell(output_start, 0), table.cell(group_end - 1, 0))
            merge_item_name_cells(table, output_start + 1, group_end)


def merge_item_name_cells(table, start: int, end: int) -> None:
    row_index = start
    while row_index < end:
        text = table.cell(row_index, 1).text.strip()
        if not text:
            row_index += 1
            continue
        merge_end = row_index
        cursor = row_index + 1
        while cursor < end and not table.cell(cursor, 1).text.strip():
            merge_end = cursor
            cursor += 1
        if merge_end > row_index:
            safe_merge(table.cell(row_index, 1), table.cell(merge_end, 1))
        row_index = cursor


def safe_merge(start_cell, end_cell) -> None:
    try:
        start_cell.merge(end_cell)
    except ValueError:
        pass


def find_row_by_first_cell(table, label: str) -> int:
    for index, row in enumerate(table.rows):
        if row.cells and row.cells[0].text.strip() == label:
            return index
    return -1


def apply_reference_table_format(table) -> None:
    column_widths = [2689, 2409, 3198]
    set_table_auto_width(table)
    set_table_grid(table, column_widths)
    for row_index, row in enumerate(table.rows):
        if 5 <= row_index <= 14:
            set_row_height(row, 113)
        for cell_index, cell in enumerate(row.cells):
            effective_col = min(cell_index, 2)
            if cell_index == 0:
                set_cell_shading(cell, "D9D9D9")
            set_cell_width(cell, column_widths[effective_col])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=0, start=108, bottom=0, end=108)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_line_spacing(paragraph, 240)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    font_name = CODE_FONT_NAME if is_code_run(run) else FONT_NAME
                    apply_run_font(run, font_name, Pt(12))
                add_line_breaks_before_inline_code(paragraph)
    apply_reference_table_xml_format(table, column_widths)


def apply_reference_table_xml_format(table, column_widths: list[int]) -> None:
    for tr in table._tbl.findall(qn("w:tr")):
        for cell_index, tc in enumerate(tr.findall(qn("w:tc"))):
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = OxmlElement("w:tcPr")
                tc.insert(0, tc_pr)
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            grid_span = tc_pr.find(qn("w:gridSpan"))
            span = int(grid_span.get(qn("w:val"), "1")) if grid_span is not None else 1
            width = sum(column_widths[cell_index : cell_index + span])
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            if cell_index == 0:
                set_tc_shading(tc, "D9D9D9")
            for p in tc.findall(qn("w:p")):
                p_pr = p.find(qn("w:pPr"))
                if p_pr is None:
                    p_pr = OxmlElement("w:pPr")
                    p.insert(0, p_pr)
                spacing = p_pr.find(qn("w:spacing"))
                if spacing is None:
                    spacing = OxmlElement("w:spacing")
                    p_pr.append(spacing)
                spacing.set(qn("w:line"), "240")
                spacing.set(qn("w:lineRule"), "auto")
                spacing.attrib.pop(qn("w:before"), None)
                spacing.attrib.pop(qn("w:after"), None)
                jc = p_pr.find(qn("w:jc"))
                if jc is None:
                    jc = OxmlElement("w:jc")
                    p_pr.append(jc)
                jc.set(qn("w:val"), "center")


def remove_empty_image_captions(doc: Document) -> None:
    # Pandoc does not create a caption for empty-alt images. This guard removes
    # any leftover duplicate paragraph from older generated files.
    seen_flowchart_heading = False
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "算法模型流程图" and paragraph.style and paragraph.style.name.startswith("Heading"):
            seen_flowchart_heading = True
            continue
        if seen_flowchart_heading and text == "算法模型流程图":
            paragraph._element.getparent().remove(paragraph._element)
            break


def add_page_break_before_first_drawing_heading(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    drawing_index = next(
        (index for index, paragraph in enumerate(paragraphs) if paragraph._p.findall(".//w:drawing", namespaces=paragraph._p.nsmap)),
        -1,
    )
    if drawing_index == -1:
        return
    for paragraph in reversed(paragraphs[:drawing_index]):
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            p_pr = paragraph._p.get_or_add_pPr()
            if p_pr.find(qn("w:pageBreakBefore")) is None:
                p_pr.append(OxmlElement("w:pageBreakBefore"))
            insert_page_break_paragraph_before(paragraph)
            return


def insert_page_break_paragraph_before(paragraph) -> None:
    previous = paragraph._p.getprevious()
    if previous is not None and previous.find(".//w:br[@w:type='page']", namespaces=paragraph._p.nsmap) is not None:
        return
    page_break_paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    page_break_paragraph.append(run)
    paragraph._p.addprevious(page_break_paragraph)


def add_page_break_after_first_drawing(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if not paragraph._p.findall(".//w:drawing", namespaces=paragraph._p.nsmap):
            continue
        insert_page_break_paragraph_after(paragraph)
        return


def insert_page_break_paragraph_after(paragraph) -> None:
    next_element = paragraph._p.getnext()
    if next_element is not None and next_element.find(".//w:br[@w:type='page']", namespaces=paragraph._p.nsmap) is not None:
        return
    page_break_paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    page_break_paragraph.append(run)
    paragraph._p.addnext(page_break_paragraph)


def center_drawing_paragraphs(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph._p.findall(".//w:drawing", namespaces=paragraph._p.nsmap):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_table_auto_width(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "0")
    tbl_w.set(qn("w:type"), "auto")


def set_table_grid(table, widths: list[int]) -> None:
    tbl = table._tbl
    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)


def set_row_height(row, height_twips: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height_twips))


def set_cell_shading(cell, fill: str) -> None:
    set_tc_shading(cell._tc, fill)


def set_tc_shading(tc, fill: str) -> None:
    tc_pr = tc.get_or_add_tcPr() if hasattr(tc, "get_or_add_tcPr") else tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        tc.insert(0, tc_pr)
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:themeFill"), "background1")
    shd.set(qn("w:themeFillShade"), "D9")


def set_paragraph_line_spacing(paragraph, line_twips: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:line"), str(line_twips))
    spacing.set(qn("w:lineRule"), "auto")


def set_table_width(table, width_twips: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_east_asia_font(rpr, font_name: str) -> None:
    if rpr is None:
        return
    r_fonts = rpr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


if __name__ == "__main__":
    main()
